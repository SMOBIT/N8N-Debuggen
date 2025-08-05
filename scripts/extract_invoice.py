import json
from docx import Document
import sys

# Usage: python3 extract_invoice.py input.docx

def extract_document(path):
    doc = Document(path)

    # 1) Header & Adresse: alle Absätze bis zur ersten Tabelle sammeln
    header_paras = []
    for para in doc.paragraphs:
        if para._element.getparent().tag.endswith('tbl'):
            break
        text = para.text.strip()
        if text:
            header_paras.append(text)

    # 2) Positionen: jede Tabelle verarbeiten
    positions = []
    for table in doc.tables:
        # Kopfzeile ermitteln
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        for row in table.rows[1:]:
            row_data = { headers[i]: row.cells[i].text.strip() for i in range(len(headers)) }
            # Nur Zeilen mit irgendeinem Wert
            if any(row_data.values()):
                positions.append(row_data)

    # 3) Bedingungen: alle Absätze nach der letzten Tabelle
    conditions = []
    last_table = doc.tables[-1]._element
    collect = False
    for para in doc.paragraphs:
        if collect and para.text.strip():
            conditions.append(para.text.strip())
        # Starte sammeln, sobald wir das Ende der letzten Tabelle sehen
        if not collect and para._element.getparent().tag.endswith('tbl') and para._element is last_table:
            collect = True

    return {
        'header': header_paras,
        'positionen': positions,
        'bedingungen': conditions
    }

if __name__ == '__main__':
    if len(sys.argv) != 2:
        print('Usage: python3 extract_invoice.py <file.docx>')
        sys.exit(1)
    data = extract_document(sys.argv[1])
    print(json.dumps(data, ensure_ascii=False, indent=2))
