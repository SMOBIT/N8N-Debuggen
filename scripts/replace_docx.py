#!/usr/bin/env python3
import sys
import json
from docx import Document

def replace_in_paragraph(paragraph, replacements):
    # jede Run einzeln bearbeiten, damit Formatierungen erhalten bleiben
    for key, val in replacements.items():
        if key in paragraph.text:
            for run in paragraph.runs:
                run.text = run.text.replace(key, val)

def replace_in_table(table, replacements):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                replace_in_paragraph(p, replacements)

def main():
    # Aufruf: replace_docx.py input.docx output.docx replacements.json
    if len(sys.argv) != 4:
        print(f"Usage: {sys.argv[0]} <input.docx> <output.docx> <replacements.json>", file=sys.stderr)
        sys.exit(1)

    in_path   = sys.argv[1]
    out_path  = sys.argv[2]
    json_path = sys.argv[3]

    # JSON mit Ersetzungen laden
    try:
        with open(json_path, 'r', encoding='utf-8') as f:
            replacements = json.load(f)
    except Exception as e:
        print(f"Error loading JSON replacements from '{json_path}': {e}", file=sys.stderr)
        sys.exit(2)

    # Dokument öffnen
    try:
        doc = Document(in_path)
    except Exception as e:
        print(f"Error opening DOCX '{in_path}': {e}", file=sys.stderr)
        sys.exit(3)

    # Absätze ersetzen
    for p in doc.paragraphs:
        replace_in_paragraph(p, replacements)
    # Tabellen ersetzen
    for tbl in doc.tables:
        replace_in_table(tbl, replacements)

    # Speichern
    try:
        doc.save(out_path)
    except Exception as e:
        print(f"Error saving DOCX to '{out_path}': {e}", file=sys.stderr)
        sys.exit(4)

if __name__ == "__main__":
    main()
